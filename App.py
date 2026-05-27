from flask import Flask, request, jsonify, send_file, Response
from flask_cors import CORS
from openai import OpenAI
from dotenv import load_dotenv
import pandas as pd
import io
import json
import time
import threading
import uuid
import os
import re

load_dotenv()

ENV_API_KEY = os.getenv("OPENAI_API_KEY", "")

app = Flask(__name__, static_folder='.', static_url_path='')
CORS(app)

# =========================
# CONFIG
# =========================

MODEL = "gpt-4.1"

SYSTEM_PROMPT = """
You are an expert SEO content writer.

Rules:
- Write professional SEO optimized content
- Do NOT mention company name or website name
- Content should be generic and usable for backlinks
- Use headings and subheadings
- Do NOT use em dashes
- Use primary keyword naturally in title
- Use primary and secondary keywords naturally in first 4 paragraphs
- Bold all target keywords
- Content should be approximately 1000 words
- Use proper formatting
- Human sounding writing
- Avoid keyword stuffing
- Make content informative and readable
"""

# In-memory job store
jobs = {}

# =========================
# SERVE FRONTEND
# =========================

@app.route('/')
def index():
    return send_file('index.html')

# =========================
# HELPERS
# =========================

def generate_article(client, primary_keyword, secondary_keyword):
    user_prompt = f"""
Write a complete SEO optimized article.

Primary Keyword:
{primary_keyword}

Secondary Keyword:
{secondary_keyword}

Instructions:
Please provide 1000 words contents for each page.
Each content should be generic on the targeted keywords without using any company or website name.
Each article needs to be SEO optimized and suitable for backlink purposes.
Each article must have an SEO optimized title having the primary keyword in it.
Use the provided target keywords naturally within the first to fourth paragraphs and make them bold.
Do not use em dashes.
Use headings and sub headings and make the content relevant and professional.
Use prepositions in keywords so that they make sense.
"""
    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt}
        ],
        temperature=0.7,
    )
    return response.choices[0].message.content


def run_generation_job(job_id, api_key, keywords_data):
    jobs[job_id]["status"] = "running"
    client = OpenAI(api_key=api_key)
    results = []

    total = len(keywords_data)
    for i, row in enumerate(keywords_data):
        if jobs[job_id].get("cancelled"):
            jobs[job_id]["status"] = "cancelled"
            return

        website = row.get("website", "")
        primary = row.get("primary_keyword", "")
        secondary = row.get("secondary_keyword", "")

        jobs[job_id]["progress"] = {
            "current": i + 1,
            "total": total,
            "current_keyword": primary
        }

        try:
            article = generate_article(client, primary, secondary)
            results.append({
                "website": website,
                "primary_keyword": primary,
                "secondary_keyword": secondary,
                "generated_article": article,
                "status": "success"
            })
        except Exception as e:
            results.append({
                "website": website,
                "primary_keyword": primary,
                "secondary_keyword": secondary,
                "generated_article": f"ERROR: {str(e)}",
                "status": "error"
            })

        time.sleep(1)

    jobs[job_id]["results"] = results
    jobs[job_id]["status"] = "completed"
    jobs[job_id]["progress"]["current"] = total


# =========================
# ROUTES
# =========================

@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({"status": "ok", "env_key_loaded": bool(ENV_API_KEY)})


@app.route("/api/generate", methods=["POST"])
def start_generation():
    data = request.get_json()
    # Frontend se aaye key ko use karo, nahi toh .env wali
    api_key = data.get("api_key", "").strip() or ENV_API_KEY
    keywords = data.get("keywords", [])

    if not api_key:
        return jsonify({"error": "API key missing — .env mein daalo ya frontend mein enter karo"}), 400
    if not keywords:
        return jsonify({"error": "keywords list is empty"}), 400

    job_id = str(uuid.uuid4())
    jobs[job_id] = {
        "status": "queued",
        "progress": {"current": 0, "total": len(keywords), "current_keyword": ""},
        "results": [],
        "cancelled": False
    }

    thread = threading.Thread(
        target=run_generation_job,
        args=(job_id, api_key, keywords),
        daemon=True
    )
    thread.start()

    return jsonify({"job_id": job_id})


@app.route("/api/job/<job_id>", methods=["GET"])
def get_job_status(job_id):
    job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    return jsonify({
        "job_id": job_id,
        "status": job["status"],
        "progress": job["progress"],
        "result_count": len(job["results"])
    })


@app.route("/api/job/<job_id>/cancel", methods=["POST"])
def cancel_job(job_id):
    job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    job["cancelled"] = True
    return jsonify({"message": "Cancellation requested"})


@app.route("/api/job/<job_id>/results", methods=["GET"])
def get_job_results(job_id):
    job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    return jsonify({
        "job_id": job_id,
        "status": job["status"],
        "results": job["results"]
    })


@app.route("/api/job/<job_id>/download/excel", methods=["GET"])
def download_excel(job_id):
    job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    if not job["results"]:
        return jsonify({"error": "No results yet"}), 400

    df = pd.DataFrame(job["results"])
    buf = io.BytesIO()
    df.to_excel(buf, index=False)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name=f"seo_articles_{job_id[:8]}.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


@app.route("/api/job/<job_id>/download/word", methods=["GET"])
def download_word(job_id):
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor

    job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    if not job["results"]:
        return jsonify({"error": "No results yet"}), 400

    results = job["results"]
    doc = DocxDocument()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)

    for i, r in enumerate(results):
        if i > 0:
            doc.add_page_break()

        primary = r.get("primary_keyword", "")
        secondary = r.get("secondary_keyword", "")
        website = r.get("website", "")
        article = r.get("generated_article", "")

        h = doc.add_heading(f'Article {i+1}: {primary}', level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        def add_meta(label, value):
            p = doc.add_paragraph()
            p.add_run(label).bold = True
            p.add_run(value)

        add_meta("Website: ", website)
        add_meta("Primary Keyword: ", primary)
        add_meta("Secondary Keyword: ", secondary)
        doc.add_paragraph()

        for line in article.split('\n'):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue
            if stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('## '):
                h = doc.add_heading(stripped[3:], level=2)
                h.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
            elif stripped.startswith('# '):
                h = doc.add_heading(stripped[2:], level=1)
                h.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            else:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', stripped)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        p.add_run(part[2:-2]).bold = True
                    else:
                        p.add_run(part)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        as_attachment=True,
        download_name=f"seo_articles_{job_id[:8]}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/api/upload-excel", methods=["POST"])
def upload_excel():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    try:
        df = pd.read_excel(file)
        df.columns = [c.strip() for c in df.columns]

        keywords = []
        for _, row in df.iterrows():
            keywords.append({
                "website": str(row.get("Website", row.get("website", ""))).strip(),
                "primary_keyword": str(row.get("Primary Keywords", row.get("primary_keyword", ""))).strip(),
                "secondary_keyword": str(row.get("Secondary Keywords", row.get("secondary_keyword", ""))).strip(),
            })
        return jsonify({"keywords": keywords, "count": len(keywords)})
    except Exception as e:
        return jsonify({"error": str(e)}), 400


if __name__ == "__main__":
    app.run(debug=True, port=5000)
